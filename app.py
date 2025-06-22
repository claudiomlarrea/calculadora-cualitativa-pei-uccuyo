# app.py

import streamlit as st
import pandas as pd
from openai import OpenAI
from io import BytesIO
from docx import Document

# âœ… Crear cliente OpenAI
client = OpenAI(api_key=st.secrets["openai"]["api_key"])

# ğŸ–¼ï¸ Logo y encabezado institucional
st.set_page_config(page_title="Calculadora Cualitativa PEI UCCuyo", page_icon="ğŸ“", layout="wide")

st.markdown("""
<style>
header {visibility: hidden;}
footer {visibility: hidden;}
.css-18e3th9 {padding-top: 1rem;}
</style>
""", unsafe_allow_html=True)

st.markdown("""
# ğŸ“ Calculadora Cualitativa PEI UCCuyo
**Universidad CatÃ³lica de Cuyo**  
SecretarÃ­a de InvestigaciÃ³n â€“ EvaluaciÃ³n institucional cualitativa del PEI  
""")

st.markdown("---")

# ğŸ“¤ Subida del archivo
uploaded_file = st.file_uploader("ğŸ“¤ Sube tu archivo Excel con actividades PEI", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    st.subheader("ğŸ“‘ Vista previa de los datos")
    st.dataframe(df)

    # SelecciÃ³n de columnas
    st.subheader("ğŸ§  Selecciona columnas con texto libre para anÃ¡lisis cualitativo")
    texto_cols = st.multiselect("Selecciona una o mÃ¡s columnas", df.columns.tolist())

    if texto_cols:
        col_joined = df[texto_cols].astype(str).agg(" ".join, axis=1)

        st.subheader("ğŸ¤– AnÃ¡lisis temÃ¡tico y de discurso por actividad")
        resultados = []

        for i, texto in enumerate(col_joined):
            prompt = f"""Analiza el siguiente texto con un enfoque cualitativo:

{texto}

Devuelve:
1. AnÃ¡lisis temÃ¡tico.
2. AnÃ¡lisis del discurso.
3. ConclusiÃ³n cualitativa."""

            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=700
                )
                resultado = response.choices[0].message.content
            except Exception as e:
                resultado = f"âŒ Error: {str(e)}"
            resultados.append(resultado)

        df["AnÃ¡lisis Cualitativo"] = resultados
        st.dataframe(df[["AnÃ¡lisis Cualitativo"]])

        # ExportaciÃ³n a Word
        def export_to_word(resultados):
            doc = Document()
            doc.add_heading("AnÃ¡lisis Cualitativo PEI", 0)
            for i, r in enumerate(resultados, 1):
                doc.add_heading(f"Actividad {i}", level=2)
                doc.add_paragraph(r)
            export_path = "/mnt/data/analisis_cualitativo_pei.docx"
            doc.save(export_path)
            return export_path

        if any("âŒ Error" not in r for r in resultados):
            docx_file = export_to_word(resultados)
            with open(docx_file, "rb") as f:
                st.download_button("ğŸ“¥ Descargar AnÃ¡lisis en Word", f, file_name="analisis_cualitativo_pei.docx")
        else:
            st.warning("âš ï¸ No se pudo generar ningÃºn anÃ¡lisis vÃ¡lido para exportar.")

        # ğŸ§  AnÃ¡lisis global del conjunto
        st.subheader("ğŸ§  AnÃ¡lisis Global del Conjunto de Actividades")

        texto_global = "\n".join(col_joined)

        prompt_global = f"""Analiza el siguiente conjunto de textos con un enfoque cualitativo integral:

{texto_global}

Devuelve un Ãºnico informe estructurado con:
1. AnÃ¡lisis temÃ¡tico general.
2. AnÃ¡lisis del discurso predominante.
3. ConclusiÃ³n cualitativa integradora."""

        try:
            respuesta_global = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt_global}],
                temperature=0.7,
                max_tokens=1000
            )
            analisis_global = respuesta_global.choices[0].message.content
            st.text_area("ğŸ“‹ Resultado del AnÃ¡lisis Global", analisis_global, height=400)

            st.download_button(
                label="ğŸ“¥ Descargar AnÃ¡lisis Global (.txt)",
                data=analisis_global,
                file_name="analisis_global_pei.txt",
                mime="text/plain"
            )
        except Exception as e:
            st.error(f"âŒ Error al generar el anÃ¡lisis global: {str(e)}")

else:
    st.info("ğŸ‘† Por favor sube un archivo Excel para comenzar.")
